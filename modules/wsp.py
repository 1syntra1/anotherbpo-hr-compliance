"""WSP / ATR report generation."""
import io
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

NAVY = "1A2744"
LIGHT_BLUE = "D9E8F5"
HEADER_BLUE = "2D9CDB"
WHITE = "FFFFFF"
GREY = "F2F2F2"


def _thin_border():
    thin = Side(style="thin")
    return Border(left=thin, right=thin, top=thin, bottom=thin)


def _hfill(hex_color):
    return PatternFill("solid", fgColor=hex_color)


def _header_row(ws, row, headers, fill_hex=HEADER_BLUE, font_color=WHITE):
    border = _thin_border()
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=col, value=h)
        c.font = Font(bold=True, color=font_color)
        c.fill = _hfill(fill_hex)
        c.alignment = Alignment(horizontal="center", wrap_text=True)
        c.border = border
    ws.row_dimensions[row].height = 28


def generate_wsp(db, company_info: dict) -> bytes:
    """Generate WSP (Workplace Skills Plan) Excel report."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "WSP"

    col_widths = [6, 20, 20, 22, 10, 12, 16, 18, 12]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.merge_cells("A1:I1")
    ws["A1"] = "WORKPLACE SKILLS PLAN (WSP)"
    ws["A1"].font = Font(bold=True, size=14, color=WHITE)
    ws["A1"].fill = _hfill(NAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 32

    info_pairs = [
        ("Company:", company_info.get("company_name", "")),
        ("SETA:", company_info.get("seta", "")),
        ("SDL Number:", company_info.get("sdl_number", "")),
        ("Skills Levy (R):", company_info.get("skills_levy", "")),
        ("Report Date:", datetime.now().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(info_pairs, 2):
        ws[f"A{i}"] = label
        ws[f"B{i}"] = value
        ws[f"A{i}"].font = Font(bold=True)
        ws[f"A{i}"].fill = _hfill(LIGHT_BLUE)

    headers = [
        "#", "Employee Name", "Employee Number", "Intervention / Programme",
        "NQF Level", "Cost (R)", "Provider", "Target Date", "Status"
    ]
    _header_row(ws, 8, headers)

    records = db.execute(
        "SELECT * FROM training_records WHERE status='Planned' ORDER BY target_date"
    ).fetchall()

    border = _thin_border()
    for i, rec in enumerate(records, 1):
        row = 8 + i
        fill = _hfill(GREY) if i % 2 == 0 else PatternFill()
        values = [
            i,
            rec["employee_name"],
            rec["employee_id"],
            rec["intervention_type"],
            rec["nqf_level"],
            rec["cost"],
            rec["provider"],
            rec["target_date"],
            rec["status"],
        ]
        for col, val in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=val)
            c.border = border
            c.fill = fill
            if col == 6 and val:
                c.number_format = "#,##0.00"
                c.alignment = Alignment(horizontal="right")

    # Total row
    total_row = 9 + len(records)
    ws.cell(row=total_row, column=1, value="TOTAL").font = Font(bold=True, color=WHITE)
    ws.cell(row=total_row, column=1).fill = _hfill(NAVY)
    if records:
        ws.cell(row=total_row, column=6,
                value=f"=SUM(F9:F{total_row-1})").font = Font(bold=True, color=WHITE)
        ws.cell(row=total_row, column=6).fill = _hfill(NAVY)
        ws.cell(row=total_row, column=6).number_format = "#,##0.00"

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def generate_atr(db, company_info: dict) -> bytes:
    """Generate ATR (Annual Training Report) Excel report."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "ATR"

    col_widths = [6, 20, 20, 22, 10, 12, 16, 14, 14, 16]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.merge_cells("A1:J1")
    ws["A1"] = "ANNUAL TRAINING REPORT (ATR)"
    ws["A1"].font = Font(bold=True, size=14, color=WHITE)
    ws["A1"].fill = _hfill(NAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 32

    info_pairs = [
        ("Company:", company_info.get("company_name", "")),
        ("SETA:", company_info.get("seta", "")),
        ("SDL Number:", company_info.get("sdl_number", "")),
        ("Report Date:", datetime.now().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(info_pairs, 2):
        ws[f"A{i}"] = label
        ws[f"B{i}"] = value
        ws[f"A{i}"].font = Font(bold=True)
        ws[f"A{i}"].fill = _hfill(LIGHT_BLUE)

    headers = [
        "#", "Employee Name", "Employee Number", "Intervention / Programme",
        "NQF Level", "Cost (R)", "Provider", "Target Date", "Actual Date", "Certificate"
    ]
    _header_row(ws, 7, headers)

    records = db.execute(
        "SELECT * FROM training_records WHERE status='Completed' ORDER BY actual_date DESC"
    ).fetchall()

    border = _thin_border()
    for i, rec in enumerate(records, 1):
        row = 7 + i
        fill = _hfill(GREY) if i % 2 == 0 else PatternFill()
        values = [
            i, rec["employee_name"], rec["employee_id"],
            rec["intervention_type"], rec["nqf_level"], rec["cost"],
            rec["provider"], rec["target_date"], rec["actual_date"], rec["certificate"]
        ]
        for col, val in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=val)
            c.border = border
            c.fill = fill
            if col == 6 and val:
                c.number_format = "#,##0.00"
                c.alignment = Alignment(horizontal="right")

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def generate_wsp_word(db, company_info: dict) -> bytes:
    """Generate WSP Word document."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed.")

    doc = Document()

    # Title
    title = doc.add_heading("WORKPLACE SKILLS PLAN (WSP)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Company info
    doc.add_heading("Company Information", level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Company Name", company_info.get("company_name", "")),
        ("SETA", company_info.get("seta", "")),
        ("SDL Number", company_info.get("sdl_number", "")),
        ("Skills Levy", f"R {company_info.get('skills_levy', 0):,.2f}"),
        ("Report Date", datetime.now().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Planned training
    doc.add_heading("Planned Training Interventions (WSP)", level=1)
    records = db.execute(
        "SELECT * FROM training_records WHERE status='Planned' ORDER BY target_date"
    ).fetchall()

    if records:
        headers = ["Employee", "Programme", "NQF", "Cost (R)", "Provider", "Target Date"]
        tbl = doc.add_table(rows=1 + len(records), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col, h in enumerate(headers):
            cell = tbl.rows[0].cells[col]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for i, rec in enumerate(records, 1):
            row = tbl.rows[i]
            row.cells[0].text = rec["employee_name"] or ""
            row.cells[1].text = rec["intervention_type"] or ""
            row.cells[2].text = rec["nqf_level"] or ""
            row.cells[3].text = f"{rec['cost']:,.2f}" if rec["cost"] else ""
            row.cells[4].text = rec["provider"] or ""
            row.cells[5].text = rec["target_date"] or ""
    else:
        doc.add_paragraph("No planned training records found.")

    doc.add_paragraph()

    # Completed training
    doc.add_heading("Completed Training (ATR)", level=1)
    completed = db.execute(
        "SELECT * FROM training_records WHERE status='Completed' ORDER BY actual_date DESC"
    ).fetchall()

    if completed:
        headers2 = ["Employee", "Programme", "NQF", "Cost (R)", "Provider", "Actual Date", "Certificate"]
        tbl2 = doc.add_table(rows=1 + len(completed), cols=len(headers2))
        tbl2.style = "Table Grid"
        for col, h in enumerate(headers2):
            cell = tbl2.rows[0].cells[col]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for i, rec in enumerate(completed, 1):
            row = tbl2.rows[i]
            row.cells[0].text = rec["employee_name"] or ""
            row.cells[1].text = rec["intervention_type"] or ""
            row.cells[2].text = rec["nqf_level"] or ""
            row.cells[3].text = f"{rec['cost']:,.2f}" if rec["cost"] else ""
            row.cells[4].text = rec["provider"] or ""
            row.cells[5].text = rec["actual_date"] or ""
            row.cells[6].text = rec["certificate"] or ""
    else:
        doc.add_paragraph("No completed training records found.")

    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated by HR Audit System — {datetime.now().strftime('%d %B %Y %H:%M')}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
